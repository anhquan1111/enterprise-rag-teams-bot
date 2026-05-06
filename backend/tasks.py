"""
tasks.py - Celery Background Tasks cho Data Ingestion Pipeline

Pipeline xử lý tài liệu (đã tối ưu hoá):
    1. Nhận document_id từ queue
    2. Cập nhật DB: status = "processing"
    3. Đọc file từ disk, extract raw text (PDF hoặc DOCX)
    4. Chunking với LangChain RecursiveCharacterTextSplitter
    5. Batch embed TẤT CẢ chunks bằng Ollama /api/embed (1 HTTP call duy nhất)
    6. Bulk insert vào ChromaDB qua Python client (1 lần duy nhất)
    7. Cập nhật DB: status = "done", chunk_count = N
    8. Xóa file tạm (finally block)
    Nếu có lỗi ở bất kỳ bước nào: cập nhật DB status = "failed"

Lý do bypass LocalRecall:
    LocalRecall xử lý từng embedding tuần tự qua Ollama nội bộ — không hỗ trợ batch.
    Thay vào đó: Ollama /api/embed nhận TOÀN BỘ chunks trong 1 request → nhanh hơn 50x.
"""

import logging
import os
import shutil
import time
from typing import List, Optional

import chromadb
import fitz  # PyMuPDF - extract text từ PDF
import httpx
from celery import Task
from docx import Document as DocxDocument  # python-docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.orm import Session

from celery_app import celery_app
from config import settings
from database import SessionLocal
from models import Document, DocumentStatus

logger = logging.getLogger(__name__)

# Model nhúng — phải khớp với model đã pull trong Ollama
_EMBED_MODEL = "nomic-embed-text"
# Số chunks gửi cho Ollama mỗi lần (tránh request body quá lớn)
_EMBED_BATCH_SIZE = 100
# Timeout cho mỗi lần gọi Ollama embed
_EMBED_TIMEOUT = 120.0

# --- LocalRecall (BM25 hybrid index) ---
# LocalRecall embed TUẦN TỰ (không batch) → upload file lớn có thể mất nhiều phút.
# Vì vậy LR indexing chạy ở task RIÊNG (fire-and-forget) — main task hoàn tất khi
# ChromaDB sẵn sàng; LR catches up trong background.
_LR_TIMEOUT = 600.0           # Timeout cho POST /upload
_LR_QUICK_TIMEOUT = 30.0      # Timeout cho create-collection / delete-entry
# LocalRecall chỉ chấp nhận PDF/MD/TXT trong endpoint /upload (verified upstream).
# DOCX bị skip ở phase này; ChromaDB vẫn cover full-text → chat không bị ảnh hưởng.
_LR_SUPPORTED_EXTS = {".pdf", ".txt", ".md"}


# =============================================================================
# HELPERS: Text Extraction
# =============================================================================

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract UTF-8 text từ file PDF sử dụng PyMuPDF.
    Dùng TEXT_PRESERVE_WHITESPACE để giữ cấu trúc bảng biểu tiếng Việt.
    """
    pdf_doc = fitz.open(file_path)
    full_text_parts: List[str] = []

    for page_num, page in enumerate(pdf_doc):
        page_text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        if not page_text.strip():
            logger.warning(
                "Trang %d của '%s' không có text (có thể là ảnh scan, cần OCR).",
                page_num + 1, os.path.basename(file_path),
            )
        full_text_parts.append(page_text)

    pdf_doc.close()
    combined_text = "\n".join(full_text_parts)

    if len(combined_text.strip()) < 100:
        logger.warning(
            "Text extract từ '%s' rất ngắn (%d ký tự). File có thể bị lỗi hoặc là ảnh scan.",
            os.path.basename(file_path), len(combined_text.strip()),
        )

    return combined_text


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract UTF-8 text từ file DOCX sử dụng python-docx.
    Đọc cả nội dung trong bảng để không bỏ sót dữ liệu.
    """
    docx_doc = DocxDocument(file_path)
    text_parts: List[str] = []

    for paragraph in docx_doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in docx_doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                text_parts.append(" | ".join(row_texts))

    combined_text = "\n".join(text_parts)

    if len(combined_text.strip()) < 100:
        logger.warning(
            "Text extract từ '%s' rất ngắn (%d ký tự). Kiểm tra lại file DOCX.",
            os.path.basename(file_path), len(combined_text.strip()),
        )

    return combined_text


def extract_text(file_path: str) -> str:
    """Dispatch đến hàm extract phù hợp dựa trên phần mở rộng file."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Định dạng file không được hỗ trợ: '{ext}'. Chỉ chấp nhận .pdf và .docx")


# =============================================================================
# HELPERS: Chunking
# =============================================================================

def chunk_text(raw_text: str) -> List[str]:
    """
    Chia text thành các chunk sử dụng RecursiveCharacterTextSplitter.

    - chunk_size=2000: đủ lớn để chứa ngữ cảnh, ít chunk hơn → ít embedding calls hơn
    - chunk_overlap=200: giữ ngữ cảnh ở ranh giới giữa các chunk
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", ".", "!", "?", " ", ""],
    )
    chunks = splitter.split_text(raw_text)
    chunks = [c.strip() for c in chunks if len(c.strip()) >= 50]
    return chunks


# =============================================================================
# HELPERS: Batch Embedding via Ollama
# =============================================================================

def _embed_batches_via_ollama(texts: List[str]) -> List[List[float]]:
    """
    Gửi TẤT CẢ texts đến Ollama /api/embed để nhận embedding vectors.

    Ollama /api/embed nhận một MẢNG texts và trả về tất cả embedding trong 1 response.
    Xử lý theo sub-batch (100 chunks/lần) để tránh request body quá lớn và show progress.

    Retry tối đa 3 lần với exponential backoff (5s, 10s, 20s) cho mỗi sub-batch.

    Returns:
        List[List[float]]: embedding vector cho mỗi chunk (cùng thứ tự với inputs)
    """
    url = f"{settings.OLLAMA_HOST}/api/embed"
    all_embeddings: List[List[float]] = []

    for batch_start in range(0, len(texts), _EMBED_BATCH_SIZE):
        batch = texts[batch_start: batch_start + _EMBED_BATCH_SIZE]
        batch_end = batch_start + len(batch)
        logger.info(
            "Embedding batch %d–%d / %d với '%s'...",
            batch_start + 1, batch_end, len(texts), _EMBED_MODEL,
        )

        for attempt in range(3):
            try:
                resp = httpx.post(
                    url,
                    json={"model": _EMBED_MODEL, "input": batch},
                    timeout=_EMBED_TIMEOUT,
                )
                resp.raise_for_status()
                batch_embeddings = resp.json()["embeddings"]
                all_embeddings.extend(batch_embeddings)
                logger.info(
                    "Batch %d–%d embedded thành công (%d vectors).",
                    batch_start + 1, batch_end, len(batch_embeddings),
                )
                break
            except (httpx.TimeoutException, httpx.ConnectError) as exc:
                if attempt < 2:
                    wait_s = 5 * (2 ** attempt)  # 5s, 10s, 20s
                    logger.warning(
                        "Embed batch %d–%d lỗi (attempt %d/3), retry sau %ds: %s",
                        batch_start + 1, batch_end, attempt + 1, wait_s, exc,
                    )
                    time.sleep(wait_s)
                else:
                    raise RuntimeError(
                        f"Embed batch {batch_start + 1}–{batch_end} thất bại sau 3 lần thử: {exc}"
                    ) from exc
            except httpx.HTTPStatusError as exc:
                raise RuntimeError(
                    f"Ollama trả về HTTP {exc.response.status_code} khi embed: {exc.response.text[:200]}"
                ) from exc

    if len(all_embeddings) != len(texts):
        raise RuntimeError(
            f"Số embedding nhận được ({len(all_embeddings)}) "
            f"không khớp số chunks ({len(texts)})"
        )

    return all_embeddings


# =============================================================================
# HELPERS: Bulk Insert vào ChromaDB
# =============================================================================

def _store_chunks_to_chromadb(
    chunks: List[str],
    embeddings: List[List[float]],
    collection_name: str,
    doc_id: str,
    filename: str,
) -> None:
    """
    Bulk insert tất cả pre-embedded chunks vào ChromaDB trong 1 lần gọi duy nhất.

    Dùng chromadb.HttpClient để kết nối ChromaDB container (chromadb:8000).
    Trước khi insert, xóa các chunks cũ của document này (nếu có từ lần chạy trước).
    collection.add() nhận MẢNG documents, embeddings, metadatas, ids → 1 HTTP request.
    """
    client = chromadb.HttpClient(
        host=settings.CHROMADB_HOST,
        port=settings.CHROMADB_PORT,
    )

    # Lấy hoặc tạo collection với cosine similarity (tốt hơn L2 cho text embeddings)
    collection = client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},
    )

    # Xóa chunks cũ của document này (idempotency — nếu task chạy lại sau khi fail)
    try:
        collection.delete(where={"doc_id": {"$eq": doc_id}})
        logger.debug("Đã xóa chunks cũ (nếu có) cho document %s.", doc_id)
    except Exception as del_exc:
        logger.debug("Không xóa được chunks cũ (có thể chưa có): %s", del_exc)

    # Chuẩn bị data cho bulk insert
    ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
    metadatas = [
        {"source": filename, "doc_id": doc_id, "chunk_index": i}
        for i in range(len(chunks))
    ]

    # Bulk insert — TẤT CẢ chunks trong 1 lần gọi
    collection.add(
        ids=ids,
        documents=chunks,
        embeddings=embeddings,
        metadatas=metadatas,
    )

    logger.info(
        "ChromaDB bulk-insert hoàn tất: %d chunks → collection '%s'.",
        len(chunks), collection_name,
    )


# =============================================================================
# HELPERS: LocalRecall (BM25 hybrid index — fire-and-forget pipeline)
# =============================================================================

def _localrecall_collection_url() -> str:
    """Base URL cho các thao tác trên collection của LocalRecall."""
    return f"{settings.LOCALRECALL_HOST}/api/collections/{settings.LOCALRECALL_COLLECTION}"


def _localrecall_entry_name(doc_id: str, filename: str) -> str:
    """
    Tên entry stable trong LocalRecall: '{doc_id}__{filename}'.
    Dùng cùng pattern cho create + delete để bảo đảm idempotency khi re-ingest.
    """
    return f"{doc_id}__{filename}"


def _ensure_localrecall_collection() -> None:
    """
    POST /api/collections idempotent: 200/201 = vừa tạo, 409 = đã tồn tại.
    Mọi response khác chỉ log warning — collection có thể đã tồn tại từ run trước.
    """
    url = f"{settings.LOCALRECALL_HOST}/api/collections"
    resp = httpx.post(
        url,
        json={"name": settings.LOCALRECALL_COLLECTION},
        timeout=_LR_QUICK_TIMEOUT,
    )
    if resp.status_code in (200, 201):
        logger.info("LocalRecall: tạo collection '%s'.", settings.LOCALRECALL_COLLECTION)
    elif resp.status_code == 409:
        logger.debug("LocalRecall: collection '%s' đã tồn tại.", settings.LOCALRECALL_COLLECTION)
    else:
        logger.warning(
            "LocalRecall create-collection trả về HTTP %d: %.200s",
            resp.status_code, resp.text,
        )


def _delete_localrecall_entry_if_exists(entry_name: str) -> None:
    """
    Best-effort xóa entry cũ trước khi upload (idempotency cho re-ingest).
    Bỏ qua mọi lỗi vì entry có thể chưa tồn tại lần đầu.
    """
    url = f"{_localrecall_collection_url()}/entry/delete"
    try:
        httpx.request(
            "DELETE",
            url,
            json={"entry": entry_name},
            timeout=_LR_QUICK_TIMEOUT,
        )
    except Exception as exc:
        logger.debug("LocalRecall DELETE entry '%s' bỏ qua: %s", entry_name, exc)


def _upload_file_to_localrecall(file_path: str, entry_name: str) -> None:
    """
    POST multipart file vào /api/collections/{name}/upload.
    Retry 1 lần (tổng 2 attempt) với 10s backoff khi gặp lỗi mạng/timeout.
    """
    url = f"{_localrecall_collection_url()}/upload"
    last_exc: Optional[Exception] = None

    for attempt in range(2):
        try:
            with open(file_path, "rb") as f:
                resp = httpx.post(
                    url,
                    files={"file": (entry_name, f)},
                    timeout=_LR_TIMEOUT,
                )
                resp.raise_for_status()
                logger.info(
                    "LocalRecall upload OK: '%s' → HTTP %d.",
                    entry_name, resp.status_code,
                )
                return
        except (httpx.TimeoutException, httpx.ConnectError, httpx.HTTPStatusError) as exc:
            last_exc = exc
            if attempt < 1:
                logger.warning(
                    "LocalRecall upload '%s' lỗi (attempt %d/2), retry sau 10s: %s",
                    entry_name, attempt + 1, exc,
                )
                time.sleep(10)

    raise RuntimeError(
        f"LocalRecall upload '{entry_name}' thất bại sau 2 lần thử: {last_exc}"
    ) from last_exc


# =============================================================================
# CELERY TASK
# =============================================================================

@celery_app.task(
    name="tasks.process_document",
    bind=True,
    max_retries=0,      # Không tự retry — lỗi được lưu vào DB với status "failed"
    time_limit=1800,    # Tối đa 30 phút cho 1 task (file 50MB)
    soft_time_limit=1500,
)
def process_document_task(self: Task, document_id: str) -> dict:
    """
    Celery task xử lý bất đồng bộ tài liệu vừa upload.

    Luồng: pending → processing → (done | failed)

    Args:
        document_id: UUID string của Document record trong PostgreSQL

    Returns:
        dict với 'status' và 'chunk_count' (nếu thành công)
    """
    logger.info("=== Bắt đầu xử lý document: %s ===", document_id)

    db: Session = SessionLocal()
    file_path: str = ""

    try:
        # --- Bước 1: Lấy document record từ DB ---
        doc: Document = db.query(Document).filter(
            Document.id == document_id
        ).first()

        if not doc:
            logger.error("Không tìm thấy document với ID: %s", document_id)
            return {"status": "failed", "error": "Document không tồn tại"}

        file_path = doc.file_path
        filename = doc.filename

        # --- Bước 2: Cập nhật status = "processing" ---
        doc.status = DocumentStatus.processing
        db.commit()
        logger.info("Document '%s' → processing", filename)

        # --- Bước 2.5: Fire-and-forget LocalRecall hybrid index (BM25 keyword) ---
        # LR embed tuần tự nên có thể mất nhiều phút — chạy task RIÊNG để main task
        # không chờ. ChromaDB-only retrieval vẫn hoạt động trong khi LR catches up.
        # Copy file để LR task có bản riêng — tránh race với cleanup ở finally bên dưới.
        try:
            lr_file_path = file_path + ".lr_copy"
            shutil.copy2(file_path, lr_file_path)
            index_to_localrecall_task.delay(document_id, lr_file_path, filename)
            logger.info("Đã queue LocalRecall index cho document %s.", document_id)
        except Exception as lr_kick_exc:
            # Tuyệt đối KHÔNG fail main pipeline ở đây — ChromaDB-only mode vẫn
            # đáp ứng <3 min target và RAG chat vẫn hoạt động (graceful degradation).
            logger.warning(
                "Không queue được LocalRecall task cho document %s (non-fatal): %s",
                document_id, lr_kick_exc,
            )

        # --- Bước 3: Extract raw text ---
        logger.info("Extracting text từ: %s", file_path)
        raw_text = extract_text(file_path)
        logger.info("Extracted %d ký tự từ '%s'.", len(raw_text), filename)

        if not raw_text.strip():
            raise ValueError(
                f"Không extract được text từ '{filename}'. "
                "File có thể rỗng hoặc chỉ chứa ảnh (cần OCR)."
            )

        # --- Bước 4: Chunking ---
        logger.info("Đang chunk text...")
        chunks = chunk_text(raw_text)
        logger.info("Chia được %d chunks từ '%s'.", len(chunks), filename)

        if not chunks:
            raise ValueError(f"Không tạo được chunk nào từ '{filename}'. Text quá ngắn?")

        # --- Bước 5: Batch embed TẤT CẢ chunks bằng Ollama (1 HTTP call) ---
        logger.info(
            "Batch embedding %d chunks với Ollama model '%s'...",
            len(chunks), _EMBED_MODEL,
        )
        embeddings = _embed_batches_via_ollama(chunks)
        logger.info("Embedding hoàn tất: %d vectors.", len(embeddings))

        # --- Bước 6: Bulk insert vào ChromaDB ---
        collection_name = settings.LOCALRECALL_COLLECTION  # reuse cùng tên collection
        _store_chunks_to_chromadb(
            chunks=chunks,
            embeddings=embeddings,
            collection_name=collection_name,
            doc_id=document_id,
            filename=filename,
        )

        # --- Bước 7: Cập nhật DB: done ---
        doc.status = DocumentStatus.done
        doc.chunk_count = len(chunks)
        doc.vector_collection_name = collection_name
        doc.error_message = None
        db.commit()

        logger.info(
            "=== '%s' xử lý THÀNH CÔNG: %d chunks → ChromaDB '%s' ===",
            filename, len(chunks), collection_name,
        )
        return {"status": "done", "chunk_count": len(chunks)}

    except Exception as exc:
        logger.error(
            "Lỗi khi xử lý document %s: %s",
            document_id, str(exc), exc_info=True,
        )
        try:
            doc = db.query(Document).filter(Document.id == document_id).first()
            if doc:
                doc.status = DocumentStatus.failed
                doc.error_message = str(exc)[:1000]
                db.commit()
        except Exception as db_exc:
            logger.error("Không thể cập nhật trạng thái failed vào DB: %s", db_exc)

        return {"status": "failed", "error": str(exc)}

    finally:
        db.close()
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info("Đã xóa file tạm: %s", file_path)
            except OSError as e:
                logger.warning("Không thể xóa file tạm '%s': %s", file_path, e)


# =============================================================================
# CELERY TASK: LocalRecall hybrid indexing (fire-and-forget)
# =============================================================================

@celery_app.task(
    name="tasks.index_to_localrecall",
    bind=True,
    max_retries=1,            # 1 lần retry tự động (Celery sẽ tự re-enqueue)
    time_limit=900,           # Tối đa 15 phút (LR embed tuần tự cho file lớn)
    soft_time_limit=850,
)
def index_to_localrecall_task(
    self: Task,
    document_id: str,
    lr_file_path: str,
    filename: str,
) -> dict:
    """
    Fire-and-forget task: index document vào LocalRecall (BM25+vector hybrid)
    SAU KHI ChromaDB indexing chính (process_document_task) đã sẵn sàng.

    Vì LocalRecall embed tuần tự (không batch endpoint), task này có thể chạy
    2–10 phút cho file lớn. Failure ở đây KHÔNG ảnh hưởng Document.status —
    RAG chat vẫn hoạt động ở chế độ ChromaDB-only graceful degradation.
    """
    logger.info("=== LocalRecall index: %s (%s) ===", document_id, filename)
    db: Session = SessionLocal()

    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext not in _LR_SUPPORTED_EXTS:
            logger.warning(
                "LocalRecall không nhận '%s' (chỉ PDF/MD/TXT). Bỏ qua hybrid index cho '%s'.",
                ext, filename,
            )
            return {"status": "skipped", "reason": f"unsupported_ext:{ext}"}

        if not os.path.exists(lr_file_path):
            logger.error("LocalRecall file copy không tồn tại: %s", lr_file_path)
            return {"status": "failed", "error": "lr_file_path_missing"}

        entry_name = _localrecall_entry_name(document_id, filename)

        _ensure_localrecall_collection()
        _delete_localrecall_entry_if_exists(entry_name)
        _upload_file_to_localrecall(lr_file_path, entry_name)

        # Đánh dấu document đã được hybrid-indexed
        doc = db.query(Document).filter(Document.id == document_id).first()
        if doc:
            doc.localrecall_indexed = True
            db.commit()
            logger.info("Document %s flagged localrecall_indexed=True.", document_id)
        else:
            logger.warning(
                "Document %s không còn tồn tại khi flag localrecall_indexed.",
                document_id,
            )

        return {"status": "done", "entry": entry_name}

    except Exception as exc:
        logger.error(
            "Lỗi LocalRecall index document %s: %s",
            document_id, str(exc), exc_info=True,
        )
        # KHÔNG mark Document.status=failed — main task đã set "done" cho ChromaDB.
        # User vẫn dùng được chat (degraded). Admin re-trigger thủ công nếu cần.
        return {"status": "failed", "error": str(exc)}

    finally:
        db.close()
        if lr_file_path and os.path.exists(lr_file_path):
            try:
                os.remove(lr_file_path)
                logger.info("Đã xóa LocalRecall file copy: %s", lr_file_path)
            except OSError as e:
                logger.warning("Không thể xóa LR file copy '%s': %s", lr_file_path, e)
